import sys
import os
import subprocess

import PyPDF2
import docx
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QDragEnterEvent, QDropEvent, QLinearGradient, QPainter, QPalette, QBrush
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QFrame, QVBoxLayout, QHBoxLayout, QMessageBox
import tempfile
from docx2pdf import convert



class DragDropFrame(QFrame):
    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.setStyleSheet("background-color: qlineargradient(x1: 0, y1: 0, x2: 1, y2: 1,"
                           "stop: 0 #1abc9c, stop: 1 #3498db);")

        self.init_ui()

    def init_ui(self):
        title_label = QLabel("文件转换", self)
        title_label.setStyleSheet("color: white; font-size: 18px; font-weight: bold;")

        drag_label = QLabel("拖放单个 Word 或 PDF 文件到指定位置", self)
        drag_label.setAlignment(Qt.AlignCenter)
        drag_label.setStyleSheet("color: white; font-size: 14px;")

        self.file_label = QLabel(" ", self)
        self.file_label.setStyleSheet("color: white; font-size: 14px;")

        vbox = QVBoxLayout()
        vbox.addStretch(1)
        vbox.addWidget(title_label, alignment=Qt.AlignCenter)
        vbox.addWidget(drag_label, alignment=Qt.AlignCenter)
        vbox.addStretch(1)
        vbox.addWidget(self.file_label, alignment=Qt.AlignCenter)
        vbox.addStretch(1)

        hbox = QHBoxLayout()
        hbox.addStretch(1)
        hbox.addLayout(vbox)
        hbox.addStretch(1)
        self.setLayout(hbox)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls() and len(event.mimeData().urls()) == 1:
            url = event.mimeData().urls()[0]
            if url.isLocalFile() and url.toLocalFile().lower().endswith(('.pdf', '.doc', '.docx')):
                self.setStyleSheet("background-color: qlineargradient(x1: 0, y1: 0, x2: 1, y2: 1,"
                                   "stop: 0 #00b894, stop: 1 #0984e3);")
                event.acceptProposedAction()
            else:
                event.ignore()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        for url in event.mimeData().urls():
            if url.isLocalFile():
                file = url.toLocalFile()
                if file.lower().endswith(('.pdf', '.doc', '.docx')):
                    self.file_label.setText(f"已拖入文件: {os.path.basename(file)}")
                    self.setStyleSheet("background-color: qlineargradient(x1: 0, y1: 0, x2: 1, y2: 1,"
                                       "stop: 0 #1abc9c, stop: 1 #3498db);")
                    self.trigger_conversion(file)

    def trigger_conversion(self, file):
        if file.lower().endswith('.pdf'):
            self.to_word(file)
        elif file.lower().endswith(('.doc', '.docx')):
            self.to_pdf(file)
        else:
            self.file_label.setText("只能转换为 Word 或 PDF 文件！")

    def to_word(self, pdf_file, docx_file="output.docx"):
        try:
            # 打开 pdf 文件
            docx_file = os.path.splitext(pdf_file)[0] + '.docx'
            with open(pdf_file, 'rb') as pdf_input:
                pdf_reader = PyPDF2.PdfReader(pdf_input)
                pdf_page = pdf_reader.pages[0]
                # 创建新的 word 文档
                doc = docx.Document()

                # 复制 pdf 内容
                pdf_page_data = pdf_page.extract_text()
                for paragraph in pdf_page_data.split("\n"):
                    doc.add_paragraph(paragraph)

            # 保存 word 文件
            (root, ext) = os.path.splitext(docx_file)
            i = 1
            while os.path.isfile(docx_file):
                docx_file = f"{root}_{i}{ext}"
                i += 1
            doc.save(docx_file)
            QMessageBox.information(self, '信息', '转换成功！')
        except Exception as e:
            QMessageBox.critical(self, '错误', '未知错误！')

    def to_pdf(self, word_file):
        try:
            # 打开 word 文件
            doc = docx.Document(word_file)
            section = doc.sections[0]
            width = section.page_width.inches * 72
            height = section.page_height.inches * 72

            # 创建空白 pdf 文件

            pdf_file = os.path.splitext(word_file)[0] + '.pdf'
            (root, ext) = os.path.splitext(pdf_file)
            i = 1
            while os.path.isfile(pdf_file):
                pdf_file = f"{root}_{i}{ext}"
                i += 1
            with PyPDF2.PdfWriter() as pdf_writer:
                pdf_page = pdf_writer.add_blank_page(width=width, height=height)
                with open(pdf_file, 'wb') as pdf_output:
                    pdf_writer.write(pdf_output)

            # 使用 docx2pdf 包将 word 文件转换为 pdf
            with tempfile.NamedTemporaryFile(delete=False) as blank_pdf:
                blank_pdf.write(b'%PDF-1.5\n')
                blank_pdf.write(b'%%EOF\n')
                convert(word_file, pdf_file, open(str(blank_pdf.name), 'rb'))

            QMessageBox.information(self, '信息', '转换成功！')
        except Exception as e:
            QMessageBox.critical(self, '错误', '未知错误！')


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("文件格式转换器")
        self.setFixedSize(400, 300)
        self.setCentralWidget(DragDropFrame())
        self.show()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    mw = MainWindow()
    sys.exit(app.exec_())
